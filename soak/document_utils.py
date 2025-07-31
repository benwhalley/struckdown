"""Text extraction utilities for PDF, Word, and text documents."""

import glob
import os
import re
import tempfile
import zipfile
from contextlib import contextmanager
from pathlib import Path

import docx
import magic
import pdfplumber
import scrubadub
import scrubadub_spacy
from pdfplumber.utils.exceptions import PdfminerException


def strip_null_bytes(obj):
    """
    Recursively strip null bytes (\\u0000) from strings in nested structures.
    """
    if isinstance(obj, str):
        return obj.replace("\u0000", "")
    elif isinstance(obj, dict):
        return {strip_null_bytes(k): strip_null_bytes(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [strip_null_bytes(i) for i in obj]
    elif isinstance(obj, tuple):
        return tuple(strip_null_bytes(i) for i in obj)
    elif isinstance(obj, set):
        return {strip_null_bytes(i) for i in obj}
    else:
        return obj


def get_scrubber(salt, model="en_core_web_md"):
    """
    Return a configured scrubber with:
    - hashed + wrapped replacements
    - spaCy NER detector (e.g. PERSON, ORG)
    - stricter email detector (avoids '@' in normal text)
    """

    from scrubadub.detectors import EmailDetector
    from scrubadub.post_processors import FilthReplacer, PrefixSuffixReplacer
    from scrubadub_spacy.detectors import SpacyEntityDetector

    class StrictEmailDetector(EmailDetector):
        """Only match proper RFC-style emails, not things like 'vague @ times'."""

        name = "strict_email"
        regex = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", re.UNICODE)

    scrubber = scrubadub.Scrubber(
        post_processor_list=[
            FilthReplacer(include_hash=True, hash_salt=salt, hash_length=4),
            PrefixSuffixReplacer(prefix="[", suffix="]"),
        ],
        detector_list=[],  # start empty to avoid default detectors
    )
    spacy_detector = SpacyEntityDetector(model=model)
    scrubber.add_detector(spacy_detector)
    scrubber.add_detector(StrictEmailDetector())

    return scrubber


def safer_extract(zip_ref, dest_dir, max_files: int = 1000):
    """
    More safely extract a zip archive to dest_dir.

    Args:
        zip_ref: zipfile.ZipFile object
        dest_dir: directory to extract to
        max_files: max number of files allowed
    """
    members = zip_ref.infolist()

    if len(members) > max_files:
        raise Exception(f"Zip contains too many files ({len(members)} > {max_files})")

    for member in members:
        # Avoid path traversal
        dest_path = os.path.abspath(os.path.join(dest_dir, member.filename))
        if not dest_path.startswith(os.path.abspath(dest_dir)):
            raise Exception(f"Unsafe path in zip: {member.filename}")

        # Block symlinks
        is_symlink = (member.external_attr >> 16) & 0o170000 == 0o120000
        if is_symlink:
            raise Exception(f"Symlink found in zip: {member.filename}")

    zip_ref.extractall(dest_dir)


import shutil
from contextlib import contextmanager


@contextmanager
def unpack_zip_to_temp_paths_if_needed(paths: list[str]) -> list[str]:
    """
    Yield list of extracted or globbed file paths. Cleans up any temp dirs.
    """
    expanded_paths = []
    temp_dirs = []

    try:
        for path in paths:
            if path.endswith(".zip") and os.path.isfile(path):
                with zipfile.ZipFile(path, "r") as zip_ref:
                    tmpdir = tempfile.mkdtemp(prefix="unpacked_zip_")
                    temp_dirs.append(tmpdir)
                    safer_extract(zip_ref, tmpdir)
                    for root, _, files in os.walk(tmpdir):
                        for f in files:
                            expanded_paths.append(os.path.join(root, f))
            else:
                expanded_paths.extend(glob.glob(path))  # expand globs

        yield expanded_paths

    finally:
        for tmpdir in temp_dirs:
            shutil.rmtree(tmpdir, ignore_errors=True)


def extract_text(path: str) -> str:
    """Extract text from various document formats.

    Supports:
    - PDF files (.pdf)
    - Word documents (.docx)
    - Plain text files (.txt, .md, etc.)

    Args:
        path: Path to the document file

    Returns:
        Extracted text content
    """
    path = Path(path)
    mtime = path.stat().st_mtime
    # print(f"Extracted {path}")
    return strip_null_bytes(_extract_text_cached(str(path), mtime))


def _extract_docx_text(path: str) -> str:
    """Extract text from a DOCX file including headers and footers."""
    try:
        doc = docx.Document(path)
        parts = []

        # Extract body text
        parts.extend(p.text for p in doc.paragraphs if p.text.strip())

        # Extract headers and footers from all sections
        for section in doc.sections:
            header = section.header
            footer = section.footer

            parts.extend(p.text for p in header.paragraphs if p.text.strip())
            parts.extend(p.text for p in footer.paragraphs if p.text.strip())

        return "\n".join(parts)

    except Exception as e:
        print(f"DOCX read failed: {path}: {e}")
        return ""


def _extract_text_cached(path: str, mtime: float) -> str:
    """Extract text based on file extension with error handling."""
    suffix = Path(path).suffix.lower()

    try:
        if suffix == ".pdf":
            with pdfplumber.open(path) as pdf:
                pages_text = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                return "\n".join(pages_text)

        elif suffix == ".docx":
            return _extract_docx_text(path)

        else:
            # Default to plain text reading
            with open(path, "r", encoding="utf-8") as f:
                return f.read()

    except PdfminerException:
        print(f"PDF read failed: {path}")
        return ""

    except Exception as e:
        print(f"Read failed: {path}: {e}")
        return ""


def is_plain_text_file(path: Path) -> bool:
    """Check if a file is a plain text file using MIME type detection."""
    try:
        mime = magic.from_file(str(path), mime=True)
        return mime.startswith("text/")
    except Exception:
        return False


def get_supported_extensions() -> list[str]:
    """Get list of supported file extensions."""
    return [".txt", ".md", ".pdf", ".docx"]


def is_supported_file(path: Path) -> bool:
    """Check if a file is supported for text extraction."""
    suffix = path.suffix.lower()
    return suffix in get_supported_extensions() or is_plain_text_file(path)


def detect_file_type(path: Path) -> str:
    """Detect the type of document for logging/debugging purposes."""
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return "PDF"
    elif suffix == ".docx":
        return "Word Document"
    elif suffix in [".txt", ".md"]:
        return "Text File"
    elif is_plain_text_file(path):
        return "Plain Text"
    else:
        return "Unknown"
